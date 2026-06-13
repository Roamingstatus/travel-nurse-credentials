"""Lightweight heuristics (no external API). Premium tier can swap in real AI parsing."""

import io
import logging as _logging
import multiprocessing as _mp
import re
import threading as _threading
from datetime import datetime
from pathlib import Path

_logger = _logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Process-isolated parser infrastructure
# ---------------------------------------------------------------------------
# All heavy parsers (pypdf, python-docx, PIL/tesseract) run in isolated worker
# processes so that:
#   1. A hard timeout can be enforced via process.terminate() / process.kill()
#      (Python threads cannot be forcibly cancelled).
#   2. Parser-bomb inputs that consume unbounded CPU/memory are killed at the
#      OS level and cannot leak resources into the parent server process.
#   3. A global semaphore caps concurrent parse processes so an attacker cannot
#      spawn an unlimited number of them through repeated requests.

# Max concurrent heavyweight parse processes across all requests.
_PARSE_BUDGET = _threading.Semaphore(4)

# Use 'fork' on Linux (faster child startup; child inherits parent modules).
# Fall back to 'spawn' on platforms where 'fork' is unavailable.
try:
    _MP_CTX = _mp.get_context("fork")
except ValueError:
    _MP_CTX = _mp.get_context("spawn")


def _bounded_parse(worker_fn, raw: bytes, timeout: float) -> str:
    """Run worker_fn(raw) in an isolated process, hard-killing it on timeout.

    Returns the worker's string result, or '' on timeout / budget-exceeded /
    any error.  The worker process is always reaped so resources are freed
    whether it finishes normally, times out, or raises an exception.
    """
    acquired = _PARSE_BUDGET.acquire(blocking=False)
    if not acquired:
        _logger.warning("[parser] Parse concurrency budget exhausted; skipping extraction.")
        return ""

    q: "_mp.Queue[str]" = _MP_CTX.Queue(maxsize=1)

    def _target() -> None:
        try:
            q.put(worker_fn(raw))
        except Exception:
            q.put("")

    proc = _MP_CTX.Process(target=_target, daemon=True)
    try:
        proc.start()
        proc.join(timeout=timeout)
        if proc.is_alive():
            _logger.warning(
                "[parser] Parse worker exceeded %.0f s; terminating process.", timeout
            )
            proc.terminate()
            proc.join(timeout=2)
            if proc.is_alive():
                proc.kill()
                proc.join(timeout=2)
            return ""
        return q.get_nowait() if not q.empty() else ""
    except Exception:
        return ""
    finally:
        if proc.is_alive():
            proc.kill()
        _PARSE_BUDGET.release()


# ---------------------------------------------------------------------------
# Module-level parse workers (must be at module scope to be picklable for
# 'spawn' context; with 'fork' closures also work, but this is safer).
# ---------------------------------------------------------------------------

def _pdf_parse_worker(raw: bytes) -> str:
    import io as _io
    from pypdf import PdfReader
    reader = PdfReader(_io.BytesIO(raw))
    if len(reader.pages) > 50:
        return ""
    parts: list[str] = []
    for page in reader.pages[:8]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            pass
    try:
        fields = reader.get_fields() or {}
        field_lines: list[str] = []
        for fn, fdata in fields.items():
            val = ""
            if hasattr(fdata, "value"):
                val = str(fdata.value or "").strip()
            elif isinstance(fdata, dict):
                val = str(fdata.get("/V") or fdata.get("value") or "").strip()
            if val and val not in ("/Off", "Off", "None"):
                field_lines.append(f"{fn}: {val}")
        if field_lines:
            parts.append("\n".join(field_lines))
    except Exception:
        pass
    return "\n".join(parts)[:15000]


def _pdf_title_worker(raw: bytes) -> str:
    import io as _io
    from pypdf import PdfReader
    reader = PdfReader(_io.BytesIO(raw))
    info = reader.metadata
    if info and hasattr(info, "title") and info.title:
        t = str(info.title).strip()
        if 3 < len(t) < 200 and not t.startswith("%"):
            return t
    return ""


def _docx_parse_worker(raw: bytes) -> str:
    import io as _io
    import zipfile
    import docx as _docx
    _MAX_UNCOMPRESSED = 50 * 1024 * 1024
    try:
        with zipfile.ZipFile(_io.BytesIO(raw)) as zf:
            if sum(info.file_size for info in zf.infolist()) > _MAX_UNCOMPRESSED:
                return ""
    except zipfile.BadZipFile:
        pass  # old binary .doc — not a ZIP; let Document() handle it
    doc = _docx.Document(_io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs[:500] if p.text.strip()]
    return "\n".join(parts)[:15000]


def _ocr_parse_worker(raw: bytes) -> str:
    import io as _io
    from PIL import Image
    import pytesseract
    _MAX_PIXELS = 4096 * 4096  # ~16 MP
    img = Image.open(_io.BytesIO(raw))
    w, h = img.size
    if w * h > _MAX_PIXELS:
        return ""
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    return pytesseract.image_to_string(img)[:15000]

# Keywords → category (first match wins; order matters)
_LICENSE_CERT_KEYS = (
    "license",
    "licence",
    "permit",
    "registration",
    "bar exam",
    "cpa",
    "notary",
    "certification",
    "certificate",
    "credential",
    "exam result",
    "pmp",
    "aws",
    "google cloud",
    "bls",
    "acls",
    "pals",
    "nrp",
    "tncc",
    "cpr",
)

_RULES: list[tuple[str, tuple[str, ...]]] = [
    ("Identity", ("passport", "drivers", "driver", "driver's", "license id", "state id", "ssn", "social security", "national id", "birth certificate", "government id")),
    ("Licenses & Certifications", _LICENSE_CERT_KEYS),
    ("Health & Compliance", ("vaccin", "immuniz", "tb ", "ppd", "drug screen", "physical exam", "osha", "hipaa", "compliance", "titer", "flu shot", "covid", "hepatitis", "varicella", "mmr", "background check")),
    ("Education", ("diploma", "transcript", "degree", "university", "college", "ged", "training completion")),
]


def infer_category(filename: str, title: str) -> str:
    blob = f"{filename} {title}".lower()
    for cat, keys in _RULES:
        if any(k in blob for k in keys):
            return cat
    return "Other"


# --- Date extraction ---

_MONTH_MAP = {
    'january': 1, 'february': 2, 'march': 3, 'april': 4,
    'may': 5, 'june': 6, 'july': 7, 'august': 8,
    'september': 9, 'october': 10, 'november': 11, 'december': 12,
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'jun': 6,
    'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
}

_MONTH_RE = r'(?:january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|oct|nov|dec)'

# Full date patterns (return day-precision datetime)
_DATE_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r'\b(\d{1,2})[/\-](\d{1,2})[/\-](20\d{2})\b'), 'mdy'),
    (re.compile(r'\b(20\d{2})[-/](\d{2})[-/](\d{2})\b'), 'ymd'),
    (re.compile(rf'\b({_MONTH_RE})[,.\s]+(\d{{1,2}})[,.\s]+(20\d{{2}})\b', re.IGNORECASE), 'mname_d_y'),
    (re.compile(rf'\b(\d{{1,2}})[,.\s]+({_MONTH_RE})[,.\s]+(20\d{{2}})\b', re.IGNORECASE), 'd_mname_y'),
]

# Month+year only patterns → produce the last day of that month
# Negative lookahead (?![/\-\d]) prevents matching inside MM/DD/YYYY
_MONTH_YEAR_PATTERNS: list[re.Pattern] = [
    re.compile(rf'\b({_MONTH_RE})[,.\s]+(20\d{{2}})\b', re.IGNORECASE),        # "March 2026"
    re.compile(r'\b(0?[1-9]|1[0-2])[/\-](20\d{2})\b'),                         # "03/2026"
    re.compile(r'\b(0?[1-9]|1[0-2])[/\-](\d{2})\b(?![/\-\d])'),               # "03/26" (MM/YY, not inside MM/DD/YYYY)
]

_EXPIRY_CTX = re.compile(
    r'(expir|expiration|exp\.?\s*(?:date|:)?|renew|renewal|'
    r'valid\s+(?:through|until|thru|to\b)|not\s+valid\s+after|'
    r'void\s+after|good\s+(?:through|until)|use\s+by|'
    r'through\s+date|thru|valid\s+thru|re.?certif|'
    r'next\s+renewal|due\s+(?:date|for\s+renewal)|'
    r'must\s+renew|expire\s*[sd]?)',
    re.IGNORECASE,
)
_ISSUE_CTX = re.compile(
    r'(issued?|issue\s*date|date\s+of\s+issue|effective|effective\s+date|'
    r'valid\s+from|start\s+date|date\s+issued|original\s+date|'
    r'begin\s+date|initial\s+date|activation|activated|granted|'
    r'license\s+date|cert(?:ified)?\s+date|awarded)',
    re.IGNORECASE,
)


def _last_day_of_month(y: int, mo: int) -> datetime:
    import calendar
    return datetime(y, mo, calendar.monthrange(y, mo)[1])


def _parse_date_match(m: re.Match, ptype: str) -> datetime | None:
    try:
        if ptype == 'mdy':
            mo, d, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        elif ptype == 'ymd':
            y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        elif ptype == 'mname_d_y':
            mo = _MONTH_MAP.get(m.group(1).lower())
            d, y = int(m.group(2)), int(m.group(3))
        elif ptype == 'd_mname_y':
            d = int(m.group(1))
            mo = _MONTH_MAP.get(m.group(2).lower())
            y = int(m.group(3))
        elif ptype == 'my_short':
            mo, yy = int(m.group(1)), int(m.group(2))
            y = 2000 + yy
            if not (1 <= mo <= 12 and 2020 <= y <= 2050):
                return None
            return _last_day_of_month(y, mo)
        else:
            return None
        if mo and 1 <= mo <= 12 and 1 <= d <= 31 and 2000 <= y <= 2050:
            return datetime(y, mo, d)
    except (ValueError, TypeError):
        pass
    return None


def _context_around(text: str, start: int, end: int, before: int = 160, after: int = 100) -> str:
    return text[max(0, start - before): end + after].lower()


def _classify_context(ctx: str) -> str | None:
    """Return 'expiry', 'issue', or None based on surrounding text."""
    if _EXPIRY_CTX.search(ctx):
        return 'expiry'
    if _ISSUE_CTX.search(ctx):
        return 'issue'
    return None


# Lines containing these patterns are copyright/boilerplate — ignore dates in them
_NOISE_LINE = re.compile(
    r'(©|\ball\s+rights\s+reserved\b|copyright|\bpublication\s+no\b|\bitem\s+no\b)',
    re.IGNORECASE,
)


def _is_noise_line(text: str, char_pos: int, lines: list[str], line_starts: list[int]) -> bool:
    """Return True if the character position falls on a boilerplate/copyright line."""
    if not lines:
        return False
    lo, hi = 0, len(line_starts) - 1
    while lo < hi:
        mid = (lo + hi + 1) // 2
        if line_starts[mid] <= char_pos:
            lo = mid
        else:
            hi = mid - 1
    return bool(_NOISE_LINE.search(lines[lo]))


def _extract_dates_from_text(text: str) -> tuple[datetime | None, datetime | None]:
    issued: datetime | None = None
    expires: datetime | None = None
    now = datetime.now()

    # Split into lines so we can check the line containing the date
    # and its immediate neighbours for label keywords.
    lines = text.splitlines()
    line_starts: list[int] = []
    pos = 0
    for ln in lines:
        line_starts.append(pos)
        pos += len(ln) + 1  # +1 for the \n

    def line_index_for(char_pos: int) -> int:
        lo, hi = 0, len(line_starts) - 1
        while lo < hi:
            mid = (lo + hi + 1) // 2
            if line_starts[mid] <= char_pos:
                lo = mid
            else:
                hi = mid - 1
        return lo

    def neighbour_ctx(char_pos: int, window: int = 2) -> str:
        li = line_index_for(char_pos)
        lo = max(0, li - window)
        hi = min(len(lines) - 1, li + window)
        return ' '.join(lines[lo:hi + 1]).lower()

    def classify(m: re.Match) -> str | None:
        raw_ctx = _context_around(text, m.start(), m.end())
        result = _classify_context(raw_ctx)
        if result:
            return result
        # Fall back to adjacent lines
        return _classify_context(neighbour_ctx(m.start()))

    # 1. Full-precision date patterns
    for pat, ptype in _DATE_PATTERNS:
        for m in pat.finditer(text):
            dt = _parse_date_match(m, ptype)
            if not dt:
                continue
            if _is_noise_line(text, m.start(), lines, line_starts):
                continue
            label = classify(m)
            if label == 'expiry':
                if expires is None:
                    expires = dt
            elif label == 'issue':
                if issued is None:
                    issued = dt
            else:
                if dt > now and expires is None:
                    expires = dt
                elif dt <= now and issued is None:
                    issued = dt

    # 2. Month+year-only patterns (lower priority — only fill gaps)
    for pat in _MONTH_YEAR_PATTERNS:
        for m in pat.finditer(text):
            if _is_noise_line(text, m.start(), lines, line_starts):
                continue
            try:
                grp = m.groups()
                if len(grp) == 2:
                    raw_mo, raw_y = grp
                    mo = _MONTH_MAP.get(str(raw_mo).lower(), None) or int(raw_mo)
                    y = int(raw_y)
                    if y < 100:          # 2-digit year → 20YY
                        y = 2000 + y
                    if not (1 <= mo <= 12 and 2020 <= y <= 2050):
                        continue
                    dt = _last_day_of_month(y, mo)
                else:
                    continue
            except (ValueError, TypeError):
                continue
            label = classify(m)
            if label == 'expiry' and expires is None:
                expires = dt
            elif label == 'issue' and issued is None:
                issued = dt
            elif label is None:
                if dt > now and expires is None:
                    expires = dt
                elif dt <= now and issued is None:
                    issued = dt

    return issued, expires


_DATE_PATTERNS_SIMPLE = (re.compile(r"(20\d{2})[-/](\d{2})[-/](\d{2})"),)


def infer_expiry_from_text(filename: str, title: str) -> datetime | None:
    """Best-effort date from filename/title (YYYY-MM-DD style)."""
    text = f"{filename} {title}"
    for pat in _DATE_PATTERNS_SIMPLE:
        m = pat.search(text)
        if not m:
            continue
        try:
            if m.lastindex == 3 and len(m.group(1)) == 4:
                y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
                return datetime(y, mo, d)
        except ValueError:
            continue
    return None


# --- Text extraction ---

def _extract_text(raw: bytes, mime_type: str, filename: str) -> str:
    mt = (mime_type or '').lower()
    fn = (filename or '').lower()
    if mt.startswith('text/') or fn.endswith(('.txt', '.csv', '.md')):
        try:
            return raw.decode('utf-8', errors='ignore')[:15000]
        except Exception:
            return ''
    if mt == 'application/pdf' or fn.endswith('.pdf'):
        return _bounded_parse(_pdf_parse_worker, raw, timeout=10)
    if (mt in ('application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword')
            or fn.endswith(('.docx', '.doc'))):
        return _bounded_parse(_docx_parse_worker, raw, timeout=15)
    if mt == 'application/rtf' or fn.endswith('.rtf'):
        try:
            import re as _re
            text = raw.decode('latin-1', errors='ignore')
            text = _re.sub(r'\\[a-z]+[-\d]*[ ]?', ' ', text)
            text = _re.sub(r'[{}\\]', '', text)
            text = _re.sub(r'\s+', ' ', text).strip()
            return text[:15000]
        except Exception:
            return ''
    if mt.startswith('image/') or fn.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.heic', '.bmp', '.tiff')):
        return _bounded_parse(_ocr_parse_worker, raw, timeout=20)
    return ''


def _extract_pdf_title(raw: bytes, mime_type: str, filename: str) -> str | None:
    mt = (mime_type or '').lower()
    fn = (filename or '').lower()
    if mt == 'application/pdf' or fn.endswith('.pdf'):
        result = _bounded_parse(_pdf_title_worker, raw, timeout=5)
        return result if result else None
    return None


def _clean_filename_as_title(filename: str) -> str:
    stem = Path(filename).stem
    stem = re.sub(r'[_\-\.]+', ' ', stem)
    stem = re.sub(r'^\d+\s*', '', stem)
    stem = stem.strip()
    return stem.title() if stem else ''


# --- Main entry point ---

def extract_document_text(raw: bytes, mime_type: str, filename: str, limit: int = 3000) -> str:
    """Public helper: return extracted plain text (up to *limit* chars).

    Used by the expiration rules engine so it can match against document
    content without re-running the full metadata pipeline.
    """
    text = _extract_text(raw, mime_type, filename)
    return (text or "")[:limit]


def extract_document_metadata(raw: bytes, mime_type: str, filename: str) -> dict:
    """
    Extract title, category, issued_at, expires_at from document content.
    Returns dict with string dates (YYYY-MM-DD) or None values.
    """
    import logging
    text = _extract_text(raw, mime_type, filename)

    title = _extract_pdf_title(raw, mime_type, filename)
    if not title:
        title = _clean_filename_as_title(filename) or None

    category = infer_category(filename, title or '')
    if category == 'Other' and text:
        category = infer_category('', text[:1000])

    issued_at, expires_at = _extract_dates_from_text(text) if text else (None, None)

    logging.warning(f"[DocExtract] issued={issued_at} expires={expires_at}")

    if not expires_at:
        expires_at = infer_expiry_from_text(filename, title or '')

    return {
        'title': title,
        'category': category if category != 'Other' else None,
        'issued_at': issued_at.strftime('%Y-%m-%d') if issued_at else None,
        'expires_at': expires_at.strftime('%Y-%m-%d') if expires_at else None,
    }
